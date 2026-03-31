#!/usr/bin/env python3
"""Parse MẪU THỬ K12 ONLINE.docx using python-docx to detect underlined answer runs."""

import sys
import os

# Fix Windows console encoding
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

def is_run_underlined(run):
    """Check if a run has underline formatting via XML <w:u/> element."""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        u_elem = rPr.find(qn('w:u'))
        if u_elem is not None:
            val = u_elem.get(qn('w:val'))
            # val can be 'single', 'double', etc. or None (means single)
            return val != 'none'
    return False

def is_run_bold(run):
    """Check if a run has bold formatting."""
    return run.bold == True

def parse_docx(filepath):
    """Parse DOCX and detect underlined runs for answer inference."""
    doc = Document(filepath)
    
    print(f"=== Parsing: {filepath} ===")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print()
    
    questions = []
    current_q = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Detect question start: "Câu N:" or "Câu N." or "Câu N (TYPE)"
        import re
        q_match = re.match(r'^Câu\s*(\d+)\s*[:.)\-]?\s*(?:\[(TF4|TF|SAQ|TL|MCQ)\]\s*)?(.*)$', text, re.IGNORECASE)
        if not q_match:
            q_match = re.match(r'^Câu\s*(\d+)\s*\((TF4|TF|SAQ|TL|MCQ)\)\s*[:.)\-]?\s*(.*)$', text, re.IGNORECASE)
        
        if q_match:
            # Save previous question
            if current_q:
                questions.append(current_q)
            
            num = int(q_match.group(1))
            hint = (q_match.group(2) or '').upper()
            q_text = q_match.group(3) or ''
            
            current_q = {
                'num': num,
                'hint': hint,
                'text': q_text,
                'runs': [],  # Store all runs with their formatting
                'options': [],
                'sub_items': [],
                'answer_line': None,
            }
            
            # Analyze runs in this paragraph
            for run in para.runs:
                if run.text.strip():
                    current_q['runs'].append({
                        'text': run.text,
                        'underlined': is_run_underlined(run),
                        'bold': is_run_bold(run),
                    })
            
            # For SAQ, extract answer from q_text if present
            if hint == 'SAQ':
                import re as re3
                # Pattern: number at end after . or ? 
                saq_match = re3.search(r'[.?]\s*(\d+[,.]?\d*)\s*$', q_text)
                if saq_match:
                    current_q['saq_answer'] = saq_match.group(1)
            continue
        
        if not current_q:
            continue
        
        # Detect options: A. B. C. D.
        opt_match = re.match(r'^([A-D])\.\s*(.+)$', text)
        if opt_match:
            letter = opt_match.group(1)
            opt_text = opt_match.group(2)
            
            # Check if letter or whole option is underlined
            runs_info = []
            for run in para.runs:
                if run.text.strip():
                    runs_info.append({
                        'text': run.text,
                        'underlined': is_run_underlined(run),
                        'bold': is_run_bold(run),
                    })
            
            current_q['options'].append({
                'letter': letter,
                'text': opt_text,
                'runs': runs_info,
            })
            continue
        
        # Detect TF4 sub-items: a) b) c) d)
        sub_match = re.match(r'^([a-d])\)\s*(.+)$', text, re.IGNORECASE)
        if sub_match:
            label = sub_match.group(1).lower()
            sub_text = sub_match.group(2)
            
            runs_info = []
            for run in para.runs:
                if run.text.strip():
                    runs_info.append({
                        'text': run.text,
                        'underlined': is_run_underlined(run),
                        'bold': is_run_bold(run),
                    })
            
            current_q['sub_items'].append({
                'label': label,
                'text': sub_text,
                'runs': runs_info,
            })
            continue
        
        # Detect answer line: "Đáp án:" or "Đáp án đúng:"
        ans_match = re.match(r'^Đáp\s*án\s*(?:đúng)?\s*[:\-]?\s*(.*)$', text, re.IGNORECASE)
        if ans_match:
            current_q['answer_line'] = ans_match.group(1).strip()
            current_q['_expect_saq_answer_next'] = True
            continue
        
        # If previous line was "Đáp án:", this line is the SAQ answer
        if current_q.get('_expect_saq_answer_next'):
            current_q['saq_answer'] = text.strip()
            del current_q['_expect_saq_answer_next']
            continue
        
        # Otherwise, add to question text
        for run in para.runs:
            if run.text.strip():
                current_q['runs'].append({
                    'text': run.text,
                    'underlined': is_run_underlined(run),
                    'bold': is_run_bold(run),
                })
        
        # For SAQ, try to extract answer from end of paragraph
        if current_q['hint'] == 'SAQ':
            # Pattern: answer is after last . or ? in the paragraph
            import re as re2
            # Look for number at end: "0,75" or "11,7" or "4"
            saq_ans_match = re2.search(r'[.?]\s*(\d+[,.]?\d*)\s*$', text)
            if saq_ans_match:
                current_q['saq_answer'] = saq_ans_match.group(1)
    
    # Save last question
    if current_q:
        questions.append(current_q)
    
    # === ANALYSIS ===
    print(f"=== Found {len(questions)} questions ===\n")
    
    for q in questions:
        print(f"--- Câu {q['num']} [{q['hint'] or 'auto'}] ---")
        
        # MCQ analysis
        if q['options']:
            print(f"  Options: {len(q['options'])}")
            for opt in q['options']:
                underlined_runs = [r for r in opt['runs'] if r['underlined']]
                if underlined_runs:
                    print(f"    {opt['letter']}. {opt['text'][:40]}... ⟵ UNDERLINED (answer={opt['letter']})")
                else:
                    print(f"    {opt['letter']}. {opt['text'][:40]}...")
        
        # TF4 analysis
        if q['sub_items']:
            print(f"  Sub-items: {len(q['sub_items'])}")
            for sub in q['sub_items']:
                # Check if label run is underlined
                label_underlined = any(r['underlined'] for r in sub['runs'] if sub['label'] in r['text'])
                if label_underlined:
                    print(f"    {sub['label']}) {sub['text'][:40]}... ⟵ LABEL UNDERLINED → ĐÚNG")
                else:
                    print(f"    {sub['label']}) {sub['text'][:40]}... ⟵ NOT underlined → SAI")
        
        # SAQ answer
        if q.get('saq_answer'):
            print(f"  SAQ answer (from text): '{q['saq_answer']}'")
        if q.get('answer_line'):
            print(f"  Answer line: '{q['answer_line']}'")
        
        # Underlined runs in question body
        ul_runs = [r for r in q['runs'] if r['underlined']]
        if ul_runs:
            print(f"  Underlined in body: {[r['text'][:30] for r in ul_runs]}")
        
        print()
    
    return questions

if __name__ == '__main__':
    # Try multiple possible filenames
    possible_files = [
        'MẪU THỬ K12 ONLINE.docx',
        'M?U TH? K12 ONLINE.docx',
        'public/Mau-Thu-K12-Online.docx',
    ]
    
    for fname in possible_files:
        if os.path.exists(fname):
            try:
                parse_docx(fname)
                break
            except Exception as e:
                print(f"Error parsing {fname}: {e}")
                import traceback
                traceback.print_exc()
    else:
        print("No DOCX file found!")
        print("Files in current dir:", os.listdir('.'))